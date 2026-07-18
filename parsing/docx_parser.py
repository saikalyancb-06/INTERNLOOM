from docx import Document
import os

def extract_text_boxes(doc):
    """
    Search document XML for <w:txbxContent> (text boxes) and extract their text.
    """
    text_box_texts = []
    try:
        body = doc.element.body
        for elem in body.iter():
            if elem.tag.endswith('txbxContent'):
                t_texts = []
                for t_elem in elem.iter():
                    if t_elem.tag.endswith('t') and t_elem.text:
                        t_texts.append(t_elem.text)
                if t_texts:
                    text_box_texts.append(" ".join(t_texts))
    except Exception as e:
        # Ignore XML walk failures, print warning or return empty
        pass
    return text_box_texts

def extract_text_from_docx(docx_path):
    """
    Parse a DOCX file, reading paragraphs, tables cell-by-cell, and text boxes.
    Returns: (text, parse_status, reason)
    """
    if not os.path.exists(docx_path):
        return "", "Failed", "File not found"
        
    try:
        doc = Document(docx_path)
        
        # 1. Paragraphs
        paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # 2. Tables (cell-by-cell)
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_cells_text = []
                for cell in row.cells:
                    cell_text = " ".join([p.text for p in cell.paragraphs if p.text.strip()])
                    if cell_text.strip():
                        row_cells_text.append(cell_text.strip())
                if row_cells_text:
                    tables_text.append(" | ".join(row_cells_text))
                    
        # 3. Text Boxes
        text_boxes_text = extract_text_boxes(doc)
        
        # Combine all parts
        all_text_sections = []
        if paragraphs_text:
            all_text_sections.append("\n".join(paragraphs_text))
        if tables_text:
            all_text_sections.append("\n--- Table Data ---\n" + "\n".join(tables_text))
        if text_boxes_text:
            all_text_sections.append("\n--- Floating Content ---\n" + "\n".join(text_boxes_text))
            
        full_text = "\n\n".join(all_text_sections).strip()
        
        if not full_text:
            return "", "Failed", "Document was empty"
            
        word_count = len(full_text.split())
        if word_count < 30:
            return full_text, "Partial", "Extremely short text extracted"
            
        return full_text, "Clean", "Parsed paragraphs, tables, and text boxes successfully"
        
    except Exception as e:
        return "", "Failed", f"DOCX parser error: {str(e)}"
